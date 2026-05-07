#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Service Timer — Controle de tempo por tipo de serviço."""

import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
import json
import urllib.request
import urllib.error
from datetime import datetime, timedelta
from pathlib import Path

try:
    from docx import Document
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ── Dimensões e cores ────────────────────────────────────────────────
W, H      = 215, 355
W_MINI    = 215
H_MINI    = 55          # altura minimizado: titlebar (28) + barra compacta (27)
BG        = "#1e1e2e"
SURFACE   = "#2a2a3e"
GREEN     = "#4ade80"
YELLOW    = "#fbbf24"
RED       = "#f87171"
BLUE      = "#60a5fa"
TEXT      = "#e2e8f0"
MUTED     = "#64748b"

DEFAULT_TIPOS = [
    "Atendimento ao Telefone",
    "Resolução de CSC",
    "Cards Deduca",
    "Resolução de Emails",
]

# ── Caminhos de dados ────────────────────────────────────────────────
DATA_DIR     = Path(__file__).parent
CONFIG_FILE  = DATA_DIR / "config.json"
RECORDS_JSON = DATA_DIR / "registros.json"
RECORDS_DOC  = DATA_DIR / "registros.docx"


# ── Helpers de arquivo ───────────────────────────────────────────────
def cfg_load():
    if CONFIG_FILE.exists():
        return json.loads(CONFIG_FILE.read_text("utf-8"))
    return {}


def cfg_save(c):
    CONFIG_FILE.write_text(json.dumps(c, ensure_ascii=False, indent=2), "utf-8")


def rec_load():
    if RECORDS_JSON.exists():
        return json.loads(RECORDS_JSON.read_text("utf-8"))
    return {"usuario": "", "registros": []}


def rec_save(d):
    RECORDS_JSON.write_text(json.dumps(d, ensure_ascii=False, indent=2), "utf-8")


# ── Aplicação ────────────────────────────────────────────────────────
class App:
    def __init__(self):
        self.cfg = cfg_load()
        if not self.cfg.get("usuario"):
            self._setup_wizard()

        self.root = tk.Tk()
        self._build_window()
        self._build_ui()
        self._reset_state()
        self.root.mainloop()

    # ── Primeiro acesso ──────────────────────────────────────────────
    def _setup_wizard(self):
        win = tk.Tk()
        win.title("Configuração — Service Timer")
        win.geometry("320x145")
        win.resizable(False, False)
        win.configure(bg=BG)

        tk.Label(win, text="Service Timer", bg=BG, fg=TEXT,
                 font=("Segoe UI", 13, "bold")).pack(pady=(18, 0))
        tk.Label(win, text="Configuração inicial — informe seu nome", bg=BG, fg=MUTED,
                 font=("Segoe UI", 9)).pack(pady=(2, 12))

        frm = tk.Frame(win, bg=BG)
        frm.pack()
        tk.Label(frm, text="Seu nome:", bg=BG, fg=TEXT,
                 font=("Segoe UI", 9)).grid(row=0, column=0, sticky="e", padx=6)
        name_v = tk.StringVar()
        tk.Entry(frm, textvariable=name_v, width=22,
                 font=("Segoe UI", 9)).grid(row=0, column=1, padx=6)

        def ok():
            n = name_v.get().strip()
            if not n:
                messagebox.showerror("Erro", "Por favor insira seu nome.", parent=win)
                return
            self.cfg["usuario"]     = n
            self.cfg["webhook_url"] = ""
            self.cfg.setdefault("tipos", DEFAULT_TIPOS)
            cfg_save(self.cfg)
            win.destroy()

        tk.Button(win, text="Salvar e continuar", command=ok,
                  bg=GREEN, fg="#000", font=("Segoe UI", 9, "bold"),
                  bd=0, padx=14, pady=6, cursor="hand2").pack(pady=14)
        win.mainloop()

    # ── Janela principal ─────────────────────────────────────────────
    def _build_window(self):
        r = self.root
        r.overrideredirect(True)
        r.attributes("-topmost", True)
        r.configure(bg=BG)
        r.geometry(f"{W}x{H}+80+80")
        r.resizable(False, False)

    def _drag_start(self, e):
        self._dx, self._dy = e.x, e.y

    def _drag_do(self, e):
        x = self.root.winfo_x() + e.x - self._dx
        y = self.root.winfo_y() + e.y - self._dy
        self.root.geometry(f"+{x}+{y}")

    # ── Construção da UI ─────────────────────────────────────────────
    def _mk_btn(self, parent, text, color, cmd, **kw):
        fg = "#000" if color in (GREEN, YELLOW) else TEXT
        return tk.Button(parent, text=text, bg=color, fg=fg,
                         font=("Segoe UI", 8, "bold"), bd=0, relief="flat",
                         activebackground=color, cursor="hand2", command=cmd, **kw)

    def _bind_drag(self, widget):
        widget.bind("<ButtonPress-1>", self._drag_start)
        widget.bind("<B1-Motion>",     self._drag_do)

    def _build_ui(self):
        # ── Titlebar (sempre visível) ────────────────────────────────
        tb = tk.Frame(self.root, bg=SURFACE, height=28)
        tb.pack(fill=tk.X)
        tb.pack_propagate(False)
        self._bind_drag(tb)

        title_lbl = tk.Label(tb, text="⏱  Service Timer", bg=SURFACE, fg=TEXT,
                             font=("Segoe UI", 9, "bold"), cursor="fleur")
        title_lbl.pack(side=tk.LEFT, padx=8)
        self._bind_drag(title_lbl)

        self._mk_btn(tb, "✕", RED, self._quit, padx=7, pady=3).pack(side=tk.RIGHT, padx=2, pady=2)
        self.btn_minimize = self._mk_btn(tb, "—", YELLOW, self._toggle_minimize, padx=7, pady=3)
        self.btn_minimize.pack(side=tk.RIGHT, padx=0, pady=2)

        # ── Barra compacta (visível só quando minimizado) ────────────
        self.compact_bar = tk.Frame(self.root, bg=SURFACE, height=27)
        # não empacotada ainda — aparece só ao minimizar
        self._bind_drag(self.compact_bar)

        ci = tk.Frame(self.compact_bar, bg=SURFACE)
        ci.pack(fill=tk.BOTH, expand=True, padx=10)
        self._bind_drag(ci)

        self.c_icon = tk.Label(ci, text="●", bg=SURFACE, fg=MUTED,
                               font=("Segoe UI", 11))
        self.c_icon.pack(side=tk.LEFT, padx=(0, 5))

        self.c_timer = tk.Label(ci, text="00:00:00", bg=SURFACE, fg=MUTED,
                                font=("Courier New", 12, "bold"))
        self.c_timer.pack(side=tk.LEFT)

        self.c_svc = tk.Label(ci, text="Aguardando", bg=SURFACE, fg=MUTED,
                              font=("Segoe UI", 8))
        self.c_svc.pack(side=tk.LEFT, padx=(7, 0))

        # ── Corpo principal ──────────────────────────────────────────
        self.body = tk.Frame(self.root, bg=BG)
        self.body.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)

        tk.Label(self.body, text=f"👤  {self.cfg['usuario']}",
                 bg=BG, fg=MUTED, font=("Segoe UI", 8)).pack(anchor="w")

        tk.Frame(self.body, bg=SURFACE, height=1).pack(fill=tk.X, pady=5)

        # Linha: "Tipo de serviço" + botão editar
        hdr = tk.Frame(self.body, bg=BG)
        hdr.pack(fill=tk.X)
        tk.Label(hdr, text="Tipo de serviço", bg=BG, fg=TEXT,
                 font=("Segoe UI", 8, "bold")).pack(side=tk.LEFT)
        self._mk_btn(hdr, "✎", SURFACE, self._edit_services,
                     padx=5, pady=1).pack(side=tk.RIGHT)

        # Combobox de serviços
        style = ttk.Style()
        style.theme_use("clam")
        style.configure("ST.TCombobox",
                        fieldbackground=SURFACE, background=SURFACE,
                        foreground=TEXT, selectbackground=SURFACE, selectforeground=TEXT)

        self.tipo_v = tk.StringVar()
        tipos = self.cfg.get("tipos", DEFAULT_TIPOS)
        self.tipo_v.set(tipos[0] if tipos else "")
        self.combo = ttk.Combobox(self.body, textvariable=self.tipo_v,
                                  values=tipos, state="readonly",
                                  width=26, style="ST.TCombobox")
        self.combo.pack(pady=(3, 6))

        # Display do timer
        disp = tk.Frame(self.body, bg=SURFACE)
        disp.pack(fill=tk.X, pady=2)
        self._bind_drag(disp)

        self.timer_lbl = tk.Label(disp, text="00:00:00", bg=SURFACE, fg=GREEN,
                                  font=("Courier New", 26, "bold"), cursor="fleur")
        self.timer_lbl.pack(pady=(8, 0))
        self._bind_drag(self.timer_lbl)

        self.status_lbl = tk.Label(disp, text="Aguardando", bg=SURFACE, fg=MUTED,
                                   font=("Segoe UI", 8))
        self.status_lbl.pack(pady=(0, 8))

        # Botões Iniciar / Pausar
        row = tk.Frame(self.body, bg=BG)
        row.pack(fill=tk.X, pady=(6, 3))

        self.btn_start = self._mk_btn(row, "▶  Iniciar", GREEN, self._start, pady=6)
        self.btn_start.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0, 2))

        self.btn_pause = self._mk_btn(row, "⏸  Pausar", YELLOW, self._toggle_pause, pady=6)
        self.btn_pause.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(2, 0))
        self.btn_pause.config(state=tk.DISABLED)

        # Botão Finalizar
        self.btn_stop = self._mk_btn(self.body, "⏹  Finalizar serviço", RED, self._stop, pady=7)
        self.btn_stop.pack(fill=tk.X, pady=3)
        self.btn_stop.config(state=tk.DISABLED)

        tk.Frame(self.body, bg=SURFACE, height=1).pack(fill=tk.X, pady=6)

        # Google Sheets
        sheets_row = tk.Frame(self.body, bg=BG)
        sheets_row.pack(fill=tk.X)
        self._mk_btn(sheets_row, "📊  Enviar ao Google Sheets", BLUE,
                     self._upload_sheets, pady=6).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0, 2))
        self._mk_btn(sheets_row, "?", SURFACE, self._show_sheets_help,
                     padx=7, pady=6).pack(side=tk.LEFT)

        # Rodapé
        bot = tk.Frame(self.body, bg=BG)
        bot.pack(fill=tk.X, pady=(6, 0))

        self.count_lbl = tk.Label(bot, text="", bg=BG, fg=MUTED, font=("Segoe UI", 7))
        self.count_lbl.pack(side=tk.LEFT)

        self._mk_btn(bot, "⚙", SURFACE, self._settings,
                     padx=6, pady=3).pack(side=tk.RIGHT)

        self._refresh_count()

    # ── Editor de serviços ───────────────────────────────────────────
    def _edit_services(self):
        win = tk.Toplevel(self.root)
        win.title("Editar Serviços")
        win.geometry("290x310")
        win.configure(bg=BG)
        win.attributes("-topmost", True)
        win.resizable(False, False)

        tk.Label(win, text="Tipos de Serviço", bg=BG, fg=TEXT,
                 font=("Segoe UI", 11, "bold")).pack(pady=(12, 6))

        # Listbox com scroll
        lf = tk.Frame(win, bg=BG)
        lf.pack(padx=12, fill=tk.BOTH, expand=True)

        sb = tk.Scrollbar(lf)
        sb.pack(side=tk.RIGHT, fill=tk.Y)

        lb = tk.Listbox(lf, yscrollcommand=sb.set,
                        bg=SURFACE, fg=TEXT,
                        selectbackground=BLUE, selectforeground="#000",
                        font=("Segoe UI", 9), bd=0, highlightthickness=0,
                        activestyle="none", height=7)
        lb.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb.config(command=lb.yview)

        for t in self.cfg.get("tipos", DEFAULT_TIPOS):
            lb.insert(tk.END, t)

        # Botões de ordenação e remoção
        bf = tk.Frame(win, bg=BG)
        bf.pack(padx=12, pady=4, fill=tk.X)

        def mover(delta):
            sel = lb.curselection()
            if not sel:
                return
            i   = sel[0]
            ni  = i + delta
            if ni < 0 or ni >= lb.size():
                return
            txt = lb.get(i)
            lb.delete(i)
            lb.insert(ni, txt)
            lb.selection_set(ni)

        def remover():
            sel = lb.curselection()
            if not sel:
                return
            if lb.size() <= 1:
                messagebox.showerror("Erro", "É necessário ao menos um serviço.", parent=win)
                return
            lb.delete(sel[0])

        tk.Button(bf, text="↑", bg=SURFACE, fg=TEXT, bd=0,
                  padx=10, pady=3, font=("Segoe UI", 9), cursor="hand2",
                  command=lambda: mover(-1)).pack(side=tk.LEFT, padx=(0, 2))
        tk.Button(bf, text="↓", bg=SURFACE, fg=TEXT, bd=0,
                  padx=10, pady=3, font=("Segoe UI", 9), cursor="hand2",
                  command=lambda: mover(1)).pack(side=tk.LEFT, padx=(0, 8))
        tk.Button(bf, text="✕ Remover", bg=RED, fg=TEXT, bd=0,
                  padx=8, pady=3, font=("Segoe UI", 8, "bold"), cursor="hand2",
                  command=remover).pack(side=tk.RIGHT)

        # Campo para adicionar novo serviço
        af = tk.Frame(win, bg=BG)
        af.pack(padx=12, pady=(0, 6), fill=tk.X)

        new_v = tk.StringVar()
        entry = tk.Entry(af, textvariable=new_v, font=("Segoe UI", 9),
                         bg=SURFACE, fg=TEXT, insertbackground=TEXT,
                         bd=0, highlightthickness=1, highlightbackground=MUTED)
        entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 4), ipady=3)

        def adicionar(event=None):
            txt = new_v.get().strip()
            if not txt:
                return
            lb.insert(tk.END, txt)
            new_v.set("")
            entry.focus()

        entry.bind("<Return>", adicionar)
        tk.Button(af, text="+ Adicionar", bg=GREEN, fg="#000", bd=0,
                  padx=8, pady=4, font=("Segoe UI", 8, "bold"),
                  cursor="hand2", command=adicionar).pack(side=tk.LEFT)

        # Salvar
        def salvar():
            novos = [lb.get(i) for i in range(lb.size())]
            if not novos:
                messagebox.showerror("Erro", "É necessário ao menos um serviço.", parent=win)
                return
            self.cfg["tipos"] = novos
            cfg_save(self.cfg)
            self.combo["values"] = novos
            if self.tipo_v.get() not in novos:
                self.tipo_v.set(novos[0])
            win.destroy()

        tk.Button(win, text="Salvar", command=salvar,
                  bg=GREEN, fg="#000", font=("Segoe UI", 9, "bold"),
                  bd=0, padx=16, pady=6, cursor="hand2").pack(pady=(4, 12))

    # ── Estado do timer ──────────────────────────────────────────────
    def _reset_state(self):
        self.running = False
        self.paused  = False
        self.t_start = None
        self.t_pause = None
        self.elapsed = timedelta()
        self.record  = None

    def _start(self):
        self.running = True
        self.paused  = False
        self.t_start = datetime.now()
        self.elapsed = timedelta()

        self.record = {
            "usuario":              self.cfg["usuario"],
            "tipo_servico":         self.tipo_v.get(),
            "data":                 self.t_start.strftime("%Y-%m-%d"),
            "inicio":               self.t_start.strftime("%H:%M:%S"),
            "pausas":               [],
            "fim":                  None,
            "tempo_total":          None,
            "tempo_total_segundos": 0,
            "enviado":              False,
        }

        self.btn_start.config(state=tk.DISABLED)
        self.btn_pause.config(state=tk.NORMAL)
        self.btn_stop.config(state=tk.NORMAL)
        self.combo.config(state=tk.DISABLED)
        self.status_lbl.config(text="▶ Em andamento", fg=GREEN)
        self.timer_lbl.config(fg=GREEN)

        # Atualiza barra compacta
        svc = self.tipo_v.get()
        self.c_icon.config(text="▶", fg=GREEN)
        self.c_timer.config(fg=GREEN)
        self.c_svc.config(text=svc[:22], fg=GREEN)
        self.compact_bar.config(bg=SURFACE)

        self._tick()

        # Minimiza automaticamente ao iniciar
        if self.body.winfo_ismapped():
            self._toggle_minimize()

    def _toggle_pause(self):
        if not self.running:
            return

        if not self.paused:
            self.paused  = True
            self.t_pause = datetime.now()
            self.elapsed += self.t_pause - self.t_start
            self.record["pausas"].append({"pausa": self.t_pause.strftime("%H:%M:%S")})

            self.btn_pause.config(text="▶  Retomar", bg=GREEN,
                                  fg="#000", activebackground=GREEN)
            self.status_lbl.config(text="⏸ Pausado", fg=YELLOW)
            self.timer_lbl.config(fg=YELLOW)

            self.c_icon.config(text="⏸", fg=YELLOW)
            self.c_timer.config(fg=YELLOW)
            self.c_svc.config(fg=YELLOW)
        else:
            self.paused  = False
            retorno      = datetime.now()
            self.t_start = retorno

            if self.record["pausas"] and "retorno" not in self.record["pausas"][-1]:
                self.record["pausas"][-1]["retorno"] = retorno.strftime("%H:%M:%S")

            self.btn_pause.config(text="⏸  Pausar", bg=YELLOW,
                                  fg="#000", activebackground=YELLOW)
            self.status_lbl.config(text="▶ Em andamento", fg=GREEN)
            self.timer_lbl.config(fg=GREEN)

            self.c_icon.config(text="▶", fg=GREEN)
            self.c_timer.config(fg=GREEN)
            self.c_svc.config(fg=GREEN)
            self._tick()

    def _stop(self):
        fim = datetime.now()
        self.running = False

        if not self.paused:
            self.elapsed += fim - self.t_start

        total = int(self.elapsed.total_seconds())

        if self.record["pausas"] and "retorno" not in self.record["pausas"][-1]:
            self.record["pausas"][-1]["retorno"] = fim.strftime("%H:%M:%S")

        self.record.update(
            fim=fim.strftime("%H:%M:%S"),
            tempo_total=str(timedelta(seconds=total)),
            tempo_total_segundos=total,
        )

        saved = dict(self.record)
        self._save(saved)
        self._reset_state()

        self.btn_start.config(state=tk.NORMAL)
        self.btn_pause.config(state=tk.DISABLED, text="⏸  Pausar",
                              bg=YELLOW, activebackground=YELLOW)
        self.btn_stop.config(state=tk.DISABLED)
        self.combo.config(state="readonly")
        self.timer_lbl.config(text="00:00:00", fg=GREEN)
        self.status_lbl.config(text="✓ Salvo!", fg=GREEN)

        # Reseta barra compacta
        self.c_icon.config(text="●", fg=MUTED)
        self.c_timer.config(text="00:00:00", fg=MUTED)
        self.c_svc.config(text="Aguardando", fg=MUTED)

        self._refresh_count()

        messagebox.showinfo(
            "Serviço finalizado",
            f"Registro salvo com sucesso!\n\n"
            f"Tipo: {saved['tipo_servico']}\n"
            f"Duração: {saved['tempo_total']}\n\n"
            f"Arquivos em:\n{DATA_DIR}",
        )

    def _tick(self):
        if not self.running or self.paused:
            return
        total = self.elapsed + (datetime.now() - self.t_start)
        s     = int(total.total_seconds())
        txt   = f"{s // 3600:02d}:{(s % 3600) // 60:02d}:{s % 60:02d}"

        self.timer_lbl.config(text=txt)
        self.c_timer.config(text=txt)   # atualiza também na barra compacta

        self.root.after(1000, self._tick)

    # ── Salvar registros ─────────────────────────────────────────────
    def _save(self, rec):
        data = rec_load()
        data["usuario"] = self.cfg["usuario"]
        data["ultima_atualizacao"] = datetime.now().isoformat(timespec="seconds")
        data["registros"].append(rec)
        rec_save(data)
        if DOCX_OK:
            self._update_doc(rec)

    def _update_doc(self, rec):
        if RECORDS_DOC.exists():
            doc = Document(str(RECORDS_DOC))
        else:
            doc = Document()
            doc.add_heading("Registro de Serviços", level=0)
            doc.add_paragraph(f"Responsável: {self.cfg['usuario']}")
            doc.add_paragraph("Gerado pelo Service Timer")
            doc.add_paragraph("─" * 50)

        doc.add_heading(f"{rec['tipo_servico']}  —  {rec['data']}", level=2)

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.rows[0].cells[0].text = "Campo"
        tbl.rows[0].cells[1].text = "Horário"

        linhas = [
            ("Início",          rec["inicio"]),
            ("Fim",             rec["fim"]),
            ("Duração total",   rec["tempo_total"]),
        ]
        for p in rec.get("pausas", []):
            linhas.append(("Pausa",            p.get("pausa",   "—")))
            linhas.append(("Retorno da pausa", p.get("retorno", "—")))

        for campo, valor in linhas:
            c = tbl.add_row().cells
            c[0].text = campo
            c[1].text = valor

        doc.add_paragraph()
        doc.save(str(RECORDS_DOC))

    # ── Google Sheets via Apps Script ────────────────────────────────
    APPS_SCRIPT = '''\
function doPost(e) {
  try {
    var dados = JSON.parse(e.postData.contents);
    var planilha = SpreadsheetApp.getActiveSpreadsheet();
    var nomeAba  = dados.usuario || "Sem nome";

    var aba = planilha.getSheetByName(nomeAba);
    if (!aba) {
      aba = planilha.insertSheet(nomeAba);
      aba.appendRow([
        "Usuário", "Data", "Tipo de Serviço",
        "Início", "Fim", "Duração", "Pausas"
      ]);
      aba.getRange(1, 1, 1, 7).setFontWeight("bold");
    }

    var registros = dados.registros || [];
    var linhas = [];
    for (var i = 0; i < registros.length; i++) {
      var r = registros[i];
      var pausas = (r.pausas || []).map(function(p) {
        return (p.pausa || "") + " → " + (p.retorno || "-");
      }).join("; ");
      linhas.push([
        r.usuario      || nomeAba,
        r.data         || "",
        r.tipo_servico || "",
        r.inicio       || "",
        r.fim          || "",
        r.tempo_total  || "",
        pausas
      ]);
    }

    if (linhas.length > 0) {
      aba.getRange(aba.getLastRow() + 1, 1, linhas.length, 7)
         .setValues(linhas);
    }

    return ContentService
      .createTextOutput(JSON.stringify({ status: "ok", linhas: linhas.length }))
      .setMimeType(ContentService.MimeType.JSON);

  } catch (err) {
    return ContentService
      .createTextOutput(JSON.stringify({ status: "erro", erro: err.message }))
      .setMimeType(ContentService.MimeType.JSON);
  }
}'''

    def _show_sheets_help(self):
        win = tk.Toplevel(self.root)
        win.title("Como configurar o Google Sheets")
        win.geometry("480x520")
        win.configure(bg=BG)
        win.attributes("-topmost", True)
        win.resizable(False, True)

        tk.Label(win, text="Configurar Google Sheets", bg=BG, fg=TEXT,
                 font=("Segoe UI", 11, "bold")).pack(pady=(14, 4))

        passos = (
            "1. Abra sua planilha no Google Sheets\n"
            "2. Clique em  Extensões > Apps Script\n"
            "3. Apague o código existente e cole o script abaixo\n"
            "4. Salve (Ctrl+S) e clique em  Implantar > Nova implantação\n"
            "5. Tipo: App da Web  |  Acesso: Qualquer pessoa\n"
            "6. Copie a URL gerada e cole em  ⚙ Config > URL do Web App"
        )
        tk.Label(win, text=passos, bg=BG, fg=TEXT,
                 font=("Segoe UI", 9), justify="left", anchor="w").pack(
                 padx=14, pady=(0, 8), fill=tk.X)

        tk.Frame(win, bg=SURFACE, height=1).pack(fill=tk.X, padx=14, pady=(0, 6))

        tk.Label(win, text="Script do Apps Script:", bg=BG, fg=MUTED,
                 font=("Segoe UI", 8, "bold")).pack(anchor="w", padx=14)

        # Caixa de texto com o script
        txt_frame = tk.Frame(win, bg=SURFACE)
        txt_frame.pack(fill=tk.BOTH, expand=True, padx=14, pady=(4, 0))

        sb = tk.Scrollbar(txt_frame)
        sb.pack(side=tk.RIGHT, fill=tk.Y)

        txt = tk.Text(txt_frame, yscrollcommand=sb.set,
                      bg="#0d1117", fg="#c9d1d9",
                      font=("Courier New", 8), bd=0,
                      wrap=tk.NONE, padx=8, pady=6)
        txt.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb.config(command=txt.yview)

        txt.insert("1.0", App.APPS_SCRIPT)
        txt.config(state=tk.DISABLED)

        # Scrollbar horizontal
        hb = tk.Scrollbar(win, orient=tk.HORIZONTAL, command=txt.xview)
        hb.pack(fill=tk.X, padx=14)
        txt.config(xscrollcommand=hb.set)

        # Botão copiar
        def copiar():
            win.clipboard_clear()
            win.clipboard_append(App.APPS_SCRIPT)
            btn_copy.config(text="✓ Copiado!")
            win.after(2000, lambda: btn_copy.config(text="📋  Copiar script"))

        btn_copy = tk.Button(win, text="📋  Copiar script", command=copiar,
                             bg=GREEN, fg="#000", font=("Segoe UI", 9, "bold"),
                             bd=0, padx=14, pady=6, cursor="hand2")
        btn_copy.pack(pady=10)

    def _upload_sheets(self):
        if not self.cfg.get("webhook_url"):
            url = simpledialog.askstring(
                "URL do Google Sheets",
                "Cole a URL do Web App (Google Apps Script):\n\n"
                "Veja instrucoes_google_sheets.txt para criar.",
                parent=self.root,
            )
            if not url:
                return
            self.cfg["webhook_url"] = url.strip()
            cfg_save(self.cfg)

        data      = rec_load()
        pendentes = [r for r in data.get("registros", []) if not r.get("enviado")]
        if not pendentes:
            messagebox.showinfo("Nada a enviar", "Todos os registros já foram enviados.")
            return

        try:
            payload = json.dumps({
                "usuario":   self.cfg["usuario"],
                "registros": pendentes,
            }).encode("utf-8")

            req = urllib.request.Request(
                self.cfg["webhook_url"],
                data=payload,
                headers={"Content-Type": "application/json"},
                method="POST",
            )

            with urllib.request.urlopen(req, timeout=15) as resp:
                resultado = json.loads(resp.read().decode("utf-8"))

            if resultado.get("status") != "ok":
                raise RuntimeError(resultado.get("erro", "Resposta inesperada do servidor."))

            for r in pendentes:
                r["enviado"] = True
            rec_save(data)

            messagebox.showinfo("Enviado!",
                                f"{resultado.get('linhas', len(pendentes))} "
                                f"registro(s) enviado(s) com sucesso!")

        except urllib.error.URLError as ex:
            messagebox.showerror("Erro de conexão", str(ex.reason))
        except Exception as ex:
            messagebox.showerror("Erro ao enviar", str(ex))

    # ── Utilitários de UI ────────────────────────────────────────────
    def _refresh_count(self):
        data      = rec_load()
        registros = data.get("registros", [])
        n = len(registros)
        p = sum(1 for r in registros if not r.get("enviado"))
        txt = f"{n} registro(s)"
        if p:
            txt += f"  •  {p} pendente(s)"
        self.count_lbl.config(text=txt)

    def _toggle_minimize(self):
        if self.body.winfo_ismapped():
            self.body.pack_forget()
            self.compact_bar.pack(fill=tk.X)
            self.root.geometry(f"{W_MINI}x{H_MINI}")
            self.btn_minimize.config(text="□")   # ícone de restaurar
        else:
            self.compact_bar.pack_forget()
            self.body.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
            self.root.geometry(f"{W}x{H}")
            self.btn_minimize.config(text="—")   # ícone de minimizar

    def _quit(self):
        if self.running:
            resposta = messagebox.askyesnocancel(
                "Fechar Service Timer",
                "Há um serviço em andamento.\n\n"
                "Deseja salvar o registro antes de fechar?",
            )
            if resposta is None:    # Cancelar — não fecha
                return
            if resposta:            # Sim — salva e fecha
                self._finalizar_e_salvar()
            # Não — fecha sem salvar
        else:
            if not messagebox.askyesno("Fechar Service Timer",
                                       "Deseja fechar o programa?"):
                return
        self.root.quit()

    def _finalizar_e_salvar(self):
        """Finaliza o serviço atual e salva sem mostrar messagebox."""
        fim = datetime.now()
        self.running = False

        if not self.paused:
            self.elapsed += fim - self.t_start

        total = int(self.elapsed.total_seconds())

        if self.record["pausas"] and "retorno" not in self.record["pausas"][-1]:
            self.record["pausas"][-1]["retorno"] = fim.strftime("%H:%M:%S")

        self.record.update(
            fim=fim.strftime("%H:%M:%S"),
            tempo_total=str(timedelta(seconds=total)),
            tempo_total_segundos=total,
        )

        self._save(dict(self.record))

    def _settings(self):
        win = tk.Toplevel(self.root)
        win.title("Configurações")
        win.geometry("360x215")
        win.configure(bg=BG)
        win.attributes("-topmost", True)
        win.resizable(False, False)

        tk.Label(win, text="Configurações", bg=BG, fg=TEXT,
                 font=("Segoe UI", 11, "bold")).pack(pady=(14, 8))

        frm = tk.Frame(win, bg=BG)
        frm.pack()

        tk.Label(frm, text="Nome do usuário:", bg=BG, fg=TEXT,
                 font=("Segoe UI", 9)).grid(row=0, column=0, sticky="e", padx=8, pady=5)
        nv = tk.StringVar(value=self.cfg.get("usuario", ""))
        tk.Entry(frm, textvariable=nv, width=24,
                 font=("Segoe UI", 9)).grid(row=0, column=1, padx=8, pady=5)

        tk.Label(frm, text="URL do Web App:", bg=BG, fg=TEXT,
                 font=("Segoe UI", 9)).grid(row=1, column=0, sticky="e", padx=8, pady=5)
        wv = tk.StringVar(value=self.cfg.get("webhook_url", ""))
        tk.Entry(frm, textvariable=wv, width=24,
                 font=("Segoe UI", 9)).grid(row=1, column=1, padx=8, pady=5)

        def save_cfg():
            self.cfg["usuario"]     = nv.get().strip()
            self.cfg["webhook_url"] = wv.get().strip()
            cfg_save(self.cfg)
            win.destroy()

        tk.Button(win, text="Salvar", command=save_cfg,
                  bg=GREEN, fg="#000", font=("Segoe UI", 9, "bold"),
                  bd=0, padx=14, pady=6, cursor="hand2").pack(pady=10)

        tk.Label(win, text=f"Pasta de dados: {DATA_DIR}",
                 bg=BG, fg=MUTED, font=("Segoe UI", 7)).pack()


if __name__ == "__main__":
    App()
